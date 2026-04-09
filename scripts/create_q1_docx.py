"""동역학 module#1 prelab 문제 1번 docx 생성 스크립트"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

def set_run_font(run, font_name='맑은 고딕', size=10, bold=False, italic=False):
    """런에 폰트 설정"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    # 한글 폰트 설정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

def set_paragraph_spacing(paragraph, line_spacing=1.0, space_after=Pt(2), space_before=Pt(0)):
    """단락 줄간격 설정"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    # 줄간격 1.0 = 240 twips
    spacing.set(qn('w:line'), str(int(240 * line_spacing)))
    spacing.set(qn('w:lineRule'), 'auto')
    if space_after is not None:
        spacing.set(qn('w:after'), str(int(space_after)))
    if space_before is not None:
        spacing.set(qn('w:before'), str(int(space_before)))

def add_text(doc, text, bold=False, size=10, align=None, space_after=Pt(2)):
    """일반 텍스트 추가"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    set_paragraph_spacing(p, line_spacing=1.0, space_after=space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

def add_mixed_text(doc, parts, space_after=Pt(2)):
    """혼합 텍스트 (bold/일반) 추가. parts: list of (text, bold, size)"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.0, space_after=space_after)
    for text, bold, size in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p

def add_equation_omml(paragraph, latex_str):
    """OMML 수식을 paragraph에 추가 (pandoc 활용)"""
    import subprocess
    import tempfile

    # pandoc으로 LaTeX -> OMML 변환
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(f'$${latex_str}$$')
        f.flush()
        tmp_md = f.name

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        tmp_docx = f.name

    try:
        subprocess.run(['pandoc', tmp_md, '-o', tmp_docx, '--mathml'],
                       capture_output=True, check=True)

        # 생성된 docx에서 수식 추출
        tmp_doc = Document(tmp_docx)
        for para in tmp_doc.paragraphs:
            for child in para._element:
                if child.tag.endswith('}oMathPara') or child.tag.endswith('}oMath'):
                    paragraph._element.append(copy.deepcopy(child))
                    return True
    except Exception as e:
        # fallback: 텍스트로 삽입
        run = paragraph.add_run(f'  [{latex_str}]')
        set_run_font(run, size=10, italic=True)
    finally:
        os.unlink(tmp_md)
        os.unlink(tmp_docx)
    return False

def add_equation_line(doc, latex_str, space_after=Pt(2)):
    """수식만 있는 줄 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.0, space_after=space_after)
    add_equation_omml(p, latex_str)
    return p

def add_text_with_equation(doc, text_before, latex_str, text_after="", bold=False, space_after=Pt(2)):
    """텍스트와 인라인 수식이 혼합된 줄"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.0, space_after=space_after)
    if text_before:
        run = p.add_run(text_before)
        set_run_font(run, size=10, bold=bold)
    if latex_str:
        add_equation_omml(p, latex_str)
    if text_after:
        run = p.add_run(text_after)
        set_run_font(run, size=10, bold=bold)
    return p

def create_report():
    doc = Document()

    # 페이지 여백 설정 (좁게)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ===== 제목 =====
    add_text(doc, '문제 1. 2D 평면 충돌에서의 충돌 각도 θ, φ 구하기', bold=True, size=11, space_after=Pt(4))

    # ===== 주어진 조건 =====
    add_text(doc, '주어진 조건', bold=True, size=10, space_after=Pt(2))
    add_text(doc, '충돌 과정에서 에너지 손실 및 운동량 변화가 없다고 가정한다 (완전 탄성 충돌).', space_after=Pt(1))

    conditions = ['3M₁ = M₂', 'v₀ = 7 m/s (충돌 전 M₁의 속도)', 'v₁ = 5 m/s (충돌 후 M₁의 속도)']
    for c in conditions:
        p = add_text(doc, f'  · {c}', space_after=Pt(0))

    # ===== 풀이 =====
    add_text(doc, '', space_after=Pt(2))  # 빈 줄
    add_text(doc, '풀이', bold=True, size=10, space_after=Pt(2))

    # (1) v2 계산
    add_text(doc, '(1) 운동에너지 보존법칙을 이용한 v₂ 계산', bold=True, size=10, space_after=Pt(1))
    add_text(doc, '완전 탄성 충돌이므로 운동에너지가 보존된다.', space_after=Pt(1))

    add_equation_line(doc, r'\frac{1}{2}M_1 v_0^2 = \frac{1}{2}M_1 v_1^2 + \frac{1}{2}M_2 v_2^2', space_after=Pt(1))

    add_text(doc, 'M₂ = 3M₁을 대입하고 양변을 ½M₁으로 나누면,', space_after=Pt(1))

    add_equation_line(doc, r'v_0^2 = v_1^2 + 3v_2^2', space_after=Pt(1))
    add_equation_line(doc, r'49 = 25 + 3v_2^2', space_after=Pt(1))
    add_equation_line(doc, r'v_2^2 = 8, \quad \therefore \ v_2 = 2\sqrt{2} \approx 2.83 \text{ m/s}', space_after=Pt(2))

    # (2) 운동량 보존법칙
    add_text(doc, '(2) 운동량 보존법칙 적용', bold=True, size=10, space_after=Pt(1))

    add_text(doc, 'x축 방향:', space_after=Pt(1))
    add_equation_line(doc, r'M_1 v_0 = M_1 v_1 \cos\phi + M_2 v_2 \cos\theta \quad \Rightarrow \quad 7 = 5\cos\phi + 6\sqrt{2}\cos\theta \quad \cdots \ (1)', space_after=Pt(1))

    add_text(doc, 'y축 방향:', space_after=Pt(1))
    add_equation_line(doc, r'0 = M_1 v_1 \sin\phi - M_2 v_2 \sin\theta \quad \Rightarrow \quad 5\sin\phi = 6\sqrt{2}\sin\theta \quad \cdots \ (2)', space_after=Pt(2))

    # (3) φ 계산
    add_text(doc, '(3) φ 계산', bold=True, size=10, space_after=Pt(1))

    add_text(doc, '식 (1)에서 6√2 cosθ = 7 − 5cosφ, 양변 제곱: 72cos²θ = (7 − 5cosφ)² ··· (3)', space_after=Pt(1))
    add_text(doc, '식 (2)에서 양변 제곱: 72sin²θ = 25sin²φ ··· (4)', space_after=Pt(1))
    add_text(doc, '식 (3) + 식 (4):', space_after=Pt(1))

    add_equation_line(doc, r'72 = (7 - 5\cos\phi)^2 + 25\sin^2\phi', space_after=Pt(1))
    add_equation_line(doc, r'= 49 - 70\cos\phi + 25(\cos^2\phi + \sin^2\phi) = 74 - 70\cos\phi', space_after=Pt(1))
    add_equation_line(doc, r'\cos\phi = \frac{2}{70} = \frac{1}{35}', space_after=Pt(1))
    add_equation_line(doc, r'\therefore \ \phi = \arccos\left(\frac{1}{35}\right) \approx 88.36°', space_after=Pt(2))

    # (4) θ 계산
    add_text(doc, '(4) θ 계산', bold=True, size=10, space_after=Pt(1))

    add_text(doc, 'cosφ = 1/35이므로, sinφ = √(1 − 1/1225) = 2√306/35', space_after=Pt(1))
    add_text(doc, '식 (2)에 대입하면:', space_after=Pt(1))

    add_equation_line(doc, r'\sin\theta = \frac{5\sin\phi}{6\sqrt{2}} = \frac{5 \cdot \frac{2\sqrt{306}}{35}}{6\sqrt{2}} = \frac{\sqrt{17}}{7}', space_after=Pt(1))
    add_equation_line(doc, r'\therefore \ \theta = \arcsin\left(\frac{\sqrt{17}}{7}\right) \approx 36.09°', space_after=Pt(2))

    # 검증
    add_text(doc, '검증', bold=True, size=10, space_after=Pt(1))

    add_text(doc, '식 (1) 대입 확인: cosθ = √(1 − 17/49) = 4√2/7 이므로,', space_after=Pt(1))
    add_equation_line(doc, r'5 \cdot \frac{1}{35} + 6\sqrt{2} \cdot \frac{4\sqrt{2}}{7} = \frac{1}{7} + \frac{48}{7} = 7 = v_0 \quad \checkmark', space_after=Pt(2))

    # 결과
    add_text(doc, '결과', bold=True, size=10, space_after=Pt(1))
    add_equation_line(doc, r'\phi = \arccos\left(\frac{1}{35}\right) \approx 88.36°, \quad \theta = \arcsin\left(\frac{\sqrt{17}}{7}\right) \approx 36.09°', space_after=Pt(2))

    # 저장
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output', '동역학_module1_prelab_Q1.docx')
    doc.save(output_path)
    print(f'저장 완료: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
